from flask import Flask, redirect, url_for, session, send_file, flash, render_template, request
from werkzeug.utils import secure_filename
from flask_sqlalchemy import SQLAlchemy
from flask_bcrypt import Bcrypt
from docx import Document
from docx.shared import Pt, RGBColor
import os

app = Flask(__name__)
app.secret_key = 'your_secret_key'

app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)
bcrypt = Bcrypt(app)

UPLOAD_FOLDER = 'static/uploads'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')

        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            flash('Username already exists. Choose a different one.', 'danger')
        else:
            new_user = User(username=username, password=hashed_password)
            db.session.add(new_user)
            db.session.commit()
            flash('Registration successful! You can now log in.', 'success')
            return redirect(url_for('login'))

    return render_template('register.html')

@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        user = User.query.filter_by(username=username).first()
        if user and bcrypt.check_password_hash(user.password, password):
            session['username'] = username
            return redirect(url_for('form'))
        else:
            flash('Invalid credentials. Please try again.', 'danger')

    return render_template('login.html')

@app.route('/form', methods=['GET', 'POST'])
def form():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        event_name = request.form['event_name']
        event_date = request.form['event_date']
        event_platform = request.form['event_platform']
        event_description = request.form['event_description']
        num_responses = request.form['num_responses']
        event_summary = request.form['event_summary']
        
        coordinators = []
        for i in range(1, 3):
            name_key = f'coordinator_name{i}'
            roll_key = f'coordinator_roll{i}'
            if name_key in request.form and roll_key in request.form:
                if request.form[name_key] and request.form[roll_key]:
                    coordinators.append(f"{i}. {request.form[name_key]} ({request.form[roll_key]})")

        if len(coordinators) < 1:
            flash("At least one coordinator is required.", "danger")
            return redirect(url_for('form'))

        images = []
        if 'images' in request.files:
            files = request.files.getlist('images')
            for file in files:
                if file and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    file.save(filepath)
                    images.append(filepath)

        doc_filename = generate_docx(event_name, event_date, event_platform, event_description, num_responses, images, event_summary, coordinators)
        return send_file(doc_filename, as_attachment=True)

    return render_template('form.html')

def generate_docx(event_name, event_date, event_platform, event_description, num_responses, images, event_summary, coordinators):
    doc = Document()
    

    default_font_size = doc.styles['Normal'].font.size or Pt(10)
    
    title = doc.add_paragraph()
    title.alignment = 1  
    run = title.add_run(f"Event name: {event_name}")
    run.bold = False
    run.font.size = Pt(default_font_size.pt * 1.5)  

    # doc.add_paragraph(f"Event Date: {event_date}")
    # doc.add_paragraph(f"Event Platform:\n{event_platform}")
    # doc.add_paragraph(f"Event Description:\n{event_description}")
    # doc.add_paragraph(f"No. of Responses: {num_responses}")
    event_date_paragraph = doc.add_paragraph()
    event_date_run = event_date_paragraph.add_run("Event Date: ")
    event_date_run.bold = True
    event_date_paragraph.add_run(event_date)

    event_platform_paragraph = doc.add_paragraph()
    event_platform_run = event_platform_paragraph.add_run("Event Platform: ")
    event_platform_run.bold = True
    event_platform_paragraph.add_run(event_platform)

    event_description_paragraph = doc.add_paragraph()
    event_description_run = event_description_paragraph.add_run("Event Description: ")
    event_description_run.bold = True
    event_description_paragraph.add_run(event_description)

    num_responses_paragraph = doc.add_paragraph()
    num_responses_run = num_responses_paragraph.add_run("No. of Responses: ")
    num_responses_run.bold = True
    num_responses_paragraph.add_run(num_responses)

    # if images:
    #     add_heading(doc, 'Uploaded Images:', level=2)
    #     for img in images:
    #         doc.add_picture(img, width=Pt(150)) 

    if images:
        add_heading(doc, 'Uploaded Images:', level=2)
        table = doc.add_table(rows=0, cols=2)
        row_cells = None

        for i, img_path in enumerate(images):
            if i % 2 == 0:
                row_cells = table.add_row().cells
            try:
                paragraph = row_cells[i % 2].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(img_path, width=Pt(200))  # adjust width as needed
            except Exception as e:
                print(f"Error inserting image: {img_path}, error: {e}")

    add_heading(doc, 'Event Completion Summary:', level=2)
    doc.add_paragraph(event_summary)
    
    add_heading(doc, 'Event Coordinators:', level=2)
    for coordinator in coordinators:
        doc.add_paragraph(coordinator)

    doc_filename = f'static/{event_name.replace(" ", "_")}_report.docx'
    doc.save(doc_filename)
    return doc_filename

def add_heading(doc, text, level):
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(12)  
    run.font.color.rgb = RGBColor(0, 0, 0) 
    return heading

@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect(url_for('login'))

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)
